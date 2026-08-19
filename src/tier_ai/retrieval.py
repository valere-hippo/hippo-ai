from __future__ import annotations

import csv
import dataclasses
import hashlib
import io
import json
import math
import re
import zipfile
from dataclasses import dataclass, field, asdict
from datetime import date, datetime
from pathlib import Path
from typing import Any, Iterable, Protocol

from .rules import infer_species_from_filename, infer_species_from_text, normalize_species_name, resolve_species_label

DEFAULT_EMBEDDING_MODEL = "BAAI/bge-m3"
DEFAULT_RERANKER_MODEL = "BAAI/bge-reranker-v2-m3"
DEFAULT_VECTOR_SIZE = 256
INDEX_FILENAME = "retrieval_index.json"

SUPPORTED_GEO_EXTENSIONS = {".gpkg", ".shp", ".geojson", ".json", ".kml", ".kmz"}
SUPPORTED_TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".tsv", ".json", ".xml", ".html", ".htm", ".qgs", ".qml", ".qgz"}
SUPPORTED_DOC_EXTENSIONS = {".pdf", ".docx", ".odt", ".rtf"}

QDRANT_AVAILABLE = False
SentenceTransformer = None
CrossEncoder = None
QdrantClient = None
Distance = None
VectorParams = None
Filter = None
FieldCondition = None
MatchValue = None
Range = None
PointStruct = None

try:  # pragma: no cover - optional dependency
    from sentence_transformers import CrossEncoder as _CrossEncoder
    from sentence_transformers import SentenceTransformer as _SentenceTransformer

    CrossEncoder = _CrossEncoder
    SentenceTransformer = _SentenceTransformer
except Exception:  # pragma: no cover - optional dependency
    SentenceTransformer = None
    CrossEncoder = None

try:  # pragma: no cover - optional dependency
    from qdrant_client import QdrantClient as _QdrantClient
    from qdrant_client.models import Distance as _Distance
    from qdrant_client.models import FieldCondition as _FieldCondition
    from qdrant_client.models import Filter as _Filter
    from qdrant_client.models import MatchValue as _MatchValue
    from qdrant_client.models import PointStruct as _PointStruct
    from qdrant_client.models import Range as _Range
    from qdrant_client.models import VectorParams as _VectorParams

    QDRANT_AVAILABLE = True
    QdrantClient = _QdrantClient
    Distance = _Distance
    VectorParams = _VectorParams
    Filter = _Filter
    FieldCondition = _FieldCondition
    MatchValue = _MatchValue
    Range = _Range
    PointStruct = _PointStruct
except Exception:  # pragma: no cover - optional dependency
    QDRANT_AVAILABLE = False

try:  # pragma: no cover - optional dependency
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover - optional dependency
    DocxDocument = None

try:  # pragma: no cover - optional dependency
    from pypdf import PdfReader
except Exception:  # pragma: no cover - optional dependency
    PdfReader = None


@dataclass(slots=True)
class RetrievalDocument:
    id: str
    project_id: str
    project_slug: str
    source_path: str
    relative_path: str
    file_name: str
    extension: str
    category: str
    title: str
    text: str
    species: str | None = None
    observed_at: str | None = None
    zone: str | None = None
    geometry_type: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)
    embedding: list[float] = field(default_factory=list)
    created_at: str = ""


@dataclass(slots=True)
class RetrievalFilter:
    species: str | None = None
    file_type: str | None = None
    category: str | None = None
    zone: str | None = None
    date_from: str | None = None
    date_to: str | None = None
    limit: int = 10


@dataclass(slots=True)
class RetrievalHit:
    id: str
    project_id: str
    project_slug: str
    score: float
    title: str
    source_path: str
    relative_path: str
    file_name: str
    extension: str
    category: str
    species: str | None = None
    observed_at: str | None = None
    zone: str | None = None
    geometry_type: str | None = None
    snippet: str = ""
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(slots=True)
class RetrievalIndexSummary:
    project_id: str
    project_slug: str
    backend: str
    indexed_documents: int
    index_path: str
    embedding_model: str
    reranker_model: str
    source_root: str
    created_at: str


@dataclass(slots=True)
class RetrievalSearchSummary:
    project_id: str
    project_slug: str
    backend: str
    query: str
    total_candidates: int
    returned_hits: int
    index_path: str
    embedding_model: str
    reranker_model: str
    hits: list[RetrievalHit] = field(default_factory=list)


class EmbeddingBackend(Protocol):
    model_name: str

    def embed(self, texts: list[str]) -> list[list[float]]:
        ...


class RerankerBackend(Protocol):
    model_name: str

    def rerank(self, query: str, documents: list[RetrievalDocument]) -> list[float]:
        ...


class HashEmbeddingBackend:
    model_name = "hash-embedding"

    def __init__(self, dimension: int = DEFAULT_VECTOR_SIZE) -> None:
        self.dimension = dimension

    def embed(self, texts: list[str]) -> list[list[float]]:
        return [self._embed_one(text) for text in texts]

    def _embed_one(self, text: str) -> list[float]:
        vector = [0.0] * self.dimension
        tokens = _tokenize(text)
        for token in tokens:
            index = int(hashlib.blake2b(token.encode("utf-8"), digest_size=8).hexdigest(), 16) % self.dimension
            vector[index] += 1.0
        return _normalize_vector(vector)


class SentenceTransformerEmbeddingBackend:
    def __init__(self, model_name: str = DEFAULT_EMBEDDING_MODEL) -> None:
        if SentenceTransformer is None:  # pragma: no cover - optional dependency
            raise RuntimeError("sentence_transformers ist nicht installiert")
        self.model_name = model_name
        self._model = SentenceTransformer(model_name)

    def embed(self, texts: list[str]) -> list[list[float]]:
        embeddings = self._model.encode(texts, normalize_embeddings=True)
        return [list(map(float, embedding)) for embedding in embeddings]


class HeuristicReranker:
    model_name = "heuristic-reranker"

    def rerank(self, query: str, documents: list[RetrievalDocument]) -> list[float]:
        query_tokens = _token_set(query)
        scores: list[float] = []
        for document in documents:
            doc_tokens = _token_set(document.text)
            overlap = len(query_tokens & doc_tokens)
            score = float(overlap)
            if document.species and _token_set(document.species) & query_tokens:
                score += 1.5
            if document.zone and _token_set(document.zone) & query_tokens:
                score += 0.8
            scores.append(score)
        return scores


class CrossEncoderReranker:
    def __init__(self, model_name: str = DEFAULT_RERANKER_MODEL) -> None:
        if CrossEncoder is None:  # pragma: no cover - optional dependency
            raise RuntimeError("sentence_transformers ist nicht installiert")
        self.model_name = model_name
        self._model = CrossEncoder(model_name)

    def rerank(self, query: str, documents: list[RetrievalDocument]) -> list[float]:
        pairs = [(query, document.text[:6000]) for document in documents]
        scores = self._model.predict(pairs)
        return [float(score) for score in scores]


class LocalProjectIndexStore:
    def __init__(self, index_root: Path) -> None:
        self.index_root = index_root
        self.index_root.mkdir(parents=True, exist_ok=True)

    def index_path(self, project_id: str) -> Path:
        return self.index_root / f"{project_id}.json"

    def load(self, project_id: str) -> tuple[dict[str, Any], list[RetrievalDocument]]:
        path = self.index_path(project_id)
        if not path.exists():
            return {}, []
        payload = json.loads(path.read_text(encoding="utf-8"))
        documents = [RetrievalDocument(**item) for item in payload.get("documents", [])]
        metadata = {key: value for key, value in payload.items() if key != "documents"}
        return metadata, documents

    def save(self, project_id: str, metadata: dict[str, Any], documents: list[RetrievalDocument]) -> None:
        payload = dict(metadata)
        payload["documents"] = [asdict(document) for document in documents]
        self.index_path(project_id).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


class QdrantProjectIndexStore:
    """Optionaler Qdrant-Speicher.

    Wenn `qdrant_client` installiert ist und ein URL-/Pfad-Target verfügbar ist,
    wird der Index zusätzlich in Qdrant gespeichert. Für lokale Entwicklung kann
    derselbe Store auch in `path=`-Modus laufen.
    """

    def __init__(self, project_id: str, collection_name: str, url: str | None = None, path: str | None = None) -> None:
        if QdrantClient is None:  # pragma: no cover - optional dependency
            raise RuntimeError("qdrant_client ist nicht installiert")
        self.project_id = project_id
        self.collection_name = collection_name
        if url:
            self.client = QdrantClient(url=url)
        elif path:
            self.client = QdrantClient(path=path)
        else:
            self.client = QdrantClient(path=str(Path("workspace") / "state" / "qdrant"))

    def ensure_collection(self, vector_size: int) -> None:
        existing = {collection.name for collection in self.client.get_collections().collections}
        if self.collection_name in existing:
            return
        self.client.create_collection(
            collection_name=self.collection_name,
            vectors_config=VectorParams(size=vector_size, distance=Distance.COSINE),
        )

    def upsert(self, documents: list[RetrievalDocument]) -> None:
        if not documents:
            return
        vector_size = len(documents[0].embedding)
        self.ensure_collection(vector_size)
        points = []
        for document in documents:
            payload = asdict(document)
            points.append(
                PointStruct(
                    id=document.id,
                    vector=document.embedding,
                    payload=payload,
                )
            )
        self.client.upsert(collection_name=self.collection_name, points=points)

    def scroll(self, filters: RetrievalFilter, limit: int) -> list[dict[str, Any]]:
        qfilter = _qdrant_filter(self.project_id, filters)
        response, _ = self.client.scroll(
            collection_name=self.collection_name,
            scroll_filter=qfilter,
            limit=limit,
            with_payload=True,
            with_vectors=False,
        )
        return [dict(point.payload or {}, score=0.0, id=str(point.id)) for point in response]

    def search(self, query_embedding: list[float], filters: RetrievalFilter, limit: int) -> list[dict[str, Any]]:
        qfilter = _qdrant_filter(self.project_id, filters)
        response = self.client.search(
            collection_name=self.collection_name,
            query_vector=query_embedding,
            query_filter=qfilter,
            limit=limit,
            with_payload=True,
        )
        return [dict(hit.payload or {}, score=float(hit.score), id=str(hit.id)) for hit in response]


def create_embedding_backend(prefer_real: bool = True) -> EmbeddingBackend:
    if prefer_real and SentenceTransformer is not None:
        try:  # pragma: no cover - optional dependency
            return SentenceTransformerEmbeddingBackend(DEFAULT_EMBEDDING_MODEL)
        except Exception:
            pass
    return HashEmbeddingBackend()


def create_reranker_backend(prefer_real: bool = True) -> RerankerBackend:
    if prefer_real and CrossEncoder is not None:
        try:  # pragma: no cover - optional dependency
            return CrossEncoderReranker(DEFAULT_RERANKER_MODEL)
        except Exception:
            pass
    return HeuristicReranker()


def index_project(
    *,
    project_id: str,
    project_slug: str,
    source_root: Path,
    index_root: Path,
    use_qdrant: bool = True,
    prefer_real_models: bool = True,
) -> RetrievalIndexSummary:
    embedding_backend = create_embedding_backend(prefer_real=prefer_real_models)
    reranker_backend = create_reranker_backend(prefer_real=prefer_real_models)
    documents = scan_project_documents(project_id=project_id, project_slug=project_slug, source_root=source_root)
    texts = [document.text for document in documents]
    embeddings = embedding_backend.embed(texts) if texts else []
    for document, embedding in zip(documents, embeddings):
        document.embedding = embedding

    created_at = now_iso()
    metadata = {
        "project_id": project_id,
        "project_slug": project_slug,
        "backend": "qdrant" if use_qdrant and QDRANT_AVAILABLE else "local",
        "embedding_model": getattr(embedding_backend, "model_name", DEFAULT_EMBEDDING_MODEL),
        "reranker_model": getattr(reranker_backend, "model_name", DEFAULT_RERANKER_MODEL),
        "source_root": str(source_root),
        "created_at": created_at,
        "qdrant_url": _qdrant_url(),
        "qdrant_path": _qdrant_path(index_root),
        "qdrant_collection": project_id,
    }

    local_store = LocalProjectIndexStore(index_root)
    local_store.save(project_id, metadata, documents)

    if use_qdrant and QDRANT_AVAILABLE:
        try:  # pragma: no cover - optional dependency
            qdrant_store = QdrantProjectIndexStore(
                project_id=project_id,
                collection_name=project_id,
                url=_qdrant_url(),
                path=_qdrant_path(index_root),
            )
            qdrant_store.upsert(documents)
            metadata["backend"] = "qdrant"
        except Exception:
            metadata["backend"] = "local"
    local_store.save(project_id, metadata, documents)

    return RetrievalIndexSummary(
        project_id=project_id,
        project_slug=project_slug,
        backend=metadata["backend"],
        indexed_documents=len(documents),
        index_path=str(local_store.index_path(project_id)),
        embedding_model=metadata["embedding_model"],
        reranker_model=metadata["reranker_model"],
        source_root=str(source_root),
        created_at=created_at,
    )


def search_project(
    *,
    project_id: str,
    project_slug: str,
    query: str,
    index_root: Path,
    filters: RetrievalFilter | None = None,
    prefer_real_models: bool = True,
) -> RetrievalSearchSummary:
    filters = filters or RetrievalFilter()
    embedding_backend = create_embedding_backend(prefer_real=prefer_real_models)
    reranker_backend = create_reranker_backend(prefer_real=prefer_real_models)
    local_store = LocalProjectIndexStore(index_root)
    metadata, documents = local_store.load(project_id)
    qdrant_backend = _load_qdrant_backend(metadata, project_id)
    resolved_species = _resolve_filter_species(filters.species)
    filters = dataclasses.replace(filters, species=resolved_species)
    filtered = [document for document in documents if _matches_filter(document, filters)]

    if not documents:
        return RetrievalSearchSummary(
            project_id=project_id,
            project_slug=project_slug,
            backend=metadata.get("backend", "local"),
            query=query,
            total_candidates=0,
            returned_hits=0,
            index_path=str(local_store.index_path(project_id)),
            embedding_model=metadata.get("embedding_model", getattr(embedding_backend, "model_name", DEFAULT_EMBEDDING_MODEL)),
            reranker_model=metadata.get("reranker_model", getattr(reranker_backend, "model_name", DEFAULT_RERANKER_MODEL)),
            hits=[],
        )

    query_text = query.strip()
    candidate_documents: list[RetrievalDocument] = []
    scored: list[tuple[float, RetrievalDocument]] = []

    if qdrant_backend is not None:
        try:
            if query_text:
                query_embedding = embedding_backend.embed([query_text])[0]
                qdrant_hits = qdrant_backend.search(query_embedding, filters, max(filters.limit * 5, 20))
            else:
                qdrant_hits = qdrant_backend.scroll(filters, max(filters.limit * 5, 20))
            candidate_documents = [_document_from_payload(hit) for hit in qdrant_hits]
            if query_text:
                rerank_scores = reranker_backend.rerank(query_text, candidate_documents) if candidate_documents else []
                reranked = sorted(zip(candidate_documents, rerank_scores), key=lambda item: item[1], reverse=True)
                scored = [(score, document) for document, score in reranked]
            else:
                scored = [(float(hit.get("score", 0.0)), document) for hit, document in zip(qdrant_hits, candidate_documents)]
        except Exception:
            candidate_documents = []
            scored = []

    if not scored:
        if query_text:
            query_embedding = embedding_backend.embed([query_text])[0]
            for document in filtered:
                score = _cosine_similarity(query_embedding, document.embedding or [])
                score += _lexical_boost(query_text, document)
                scored.append((score, document))
            scored.sort(key=lambda item: item[0], reverse=True)
            candidate_documents = [document for _, document in scored[: max(filters.limit, 20)]]
            rerank_scores = reranker_backend.rerank(query_text, candidate_documents) if candidate_documents else []
            reranked = [(score, document) for document, score in zip(candidate_documents, rerank_scores)]
            if reranked:
                reranked.sort(key=lambda item: item[0], reverse=True)
                scored = reranked
        else:
            scored = [(0.0, document) for document in filtered]

    reranked = scored

    hits = [
        RetrievalHit(
            id=document.id,
            project_id=document.project_id,
            project_slug=document.project_slug,
            score=float(score),
            title=document.title,
            source_path=document.source_path,
            relative_path=document.relative_path,
            file_name=document.file_name,
            extension=document.extension,
            category=document.category,
            species=document.species,
            observed_at=document.observed_at,
            zone=document.zone,
            geometry_type=document.geometry_type,
            snippet=_make_snippet(document.text, query),
            metadata=document.metadata,
        )
        for score, document in reranked[: filters.limit]
    ]

    return RetrievalSearchSummary(
        project_id=project_id,
        project_slug=project_slug,
        backend=metadata.get("backend", "local"),
        query=query,
        total_candidates=len(filtered if filtered else candidate_documents),
        returned_hits=len(hits),
        index_path=str(local_store.index_path(project_id)),
        embedding_model=metadata.get("embedding_model", getattr(embedding_backend, "model_name", DEFAULT_EMBEDDING_MODEL)),
        reranker_model=metadata.get("reranker_model", getattr(reranker_backend, "model_name", DEFAULT_RERANKER_MODEL)),
        hits=hits,
    )


def scan_project_documents(*, project_id: str, project_slug: str, source_root: Path) -> list[RetrievalDocument]:
    documents: list[RetrievalDocument] = []
    for path in iter_project_files(source_root):
        category = classify_extension(path.suffix)
        if category == "geodata":
            documents.extend(_scan_geodata_file(project_id, project_slug, source_root, path))
        elif category == "document":
            documents.extend(_scan_document_file(project_id, project_slug, source_root, path))
        else:
            documents.append(_build_file_document(project_id, project_slug, source_root, path, category))
    return documents


def iter_project_files(root: Path) -> Iterable[Path]:
    excluded_dirs = {".git", ".venv", "__pycache__", "node_modules", "workspace"}
    for path in root.rglob("*"):
        if any(part in excluded_dirs for part in path.parts):
            continue
        if not path.is_file():
            continue
        if path.name in {"manifest.json", "inventory.json"}:
            continue
        yield path


def classify_extension(extension: str) -> str:
    ext = extension.lower().lstrip(".")
    if ext in {"gpkg", "shp", "geojson", "json", "kml", "kmz", "csv", "gdb", "tif", "tiff"}:
        return "geodata"
    if ext in {"qgs", "qgz"}:
        return "qgis"
    if ext in {"doc", "docx", "pdf", "rtf", "odt", "ppt", "pptx"}:
        return "document"
    if ext in {"txt", "md", "xml", "html", "htm", "csv", "tsv"}:
        return "document"
    if ext in {"png", "jpg", "jpeg", "webp", "svg"}:
        return "image"
    return "other"


def _scan_geodata_file(project_id: str, project_slug: str, source_root: Path, path: Path) -> list[RetrievalDocument]:
    docs = _scan_geojson_like(project_id, project_slug, source_root, path)
    if docs:
        return docs

    try:  # pragma: no cover - optional dependency
        import geopandas as gpd
    except Exception:
        return [_build_file_document(project_id, project_slug, source_root, path, "geodata")]

    try:  # pragma: no cover - optional dependency
        frame = gpd.read_file(path)
    except Exception:
        return [_build_file_document(project_id, project_slug, source_root, path, "geodata")]

    if frame.empty:
        return [_build_file_document(project_id, project_slug, source_root, path, "geodata")]

    species_column = _find_column(frame.columns, ["species", "art", "artname", "taxon", "taxon_name", "name", "wissenschaftlicher_name", "deutscher_name", "objektart", "bezeichnung"])
    date_column = _find_column(frame.columns, ["observed_at", "date", "datum", "beobachtet_am"])

    docs = []
    for index, row in frame.iterrows():
        geometry = row.geometry
        if geometry is None or getattr(geometry, "is_empty", False):
            continue
        attrs = {str(key): value for key, value in row.drop(labels=["geometry"], errors="ignore").to_dict().items()}
        species = _resolve_species_from_attrs(attrs, species_column)
        observed_at = _normalize_date_value(attrs.get(date_column)) if date_column else None
        zone = _infer_zone(attrs)
        geometry_type = str(getattr(geometry, "geom_type", "unknown"))
        title = species or path.stem
        text = _compose_feature_text(path, attrs, species, observed_at, zone, geometry_type)
        doc = _build_document(
            project_id=project_id,
            project_slug=project_slug,
            source_root=source_root,
            path=path,
            category="geodata",
            title=title,
            text=text,
            species=species,
            observed_at=observed_at,
            zone=zone,
            geometry_type=geometry_type,
            metadata={
                "row_index": int(index),
                "attribute_keys": sorted(attrs.keys()),
            },
        )
        docs.append(doc)

    if not docs:
        docs.append(_build_file_document(project_id, project_slug, source_root, path, "geodata"))
    return docs


def _scan_geojson_like(project_id: str, project_slug: str, source_root: Path, path: Path) -> list[RetrievalDocument]:
    if path.suffix.lower() not in {".geojson", ".json"}:
        return []
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return []
    if not isinstance(payload, dict) or payload.get("type") != "FeatureCollection":
        return []

    docs: list[RetrievalDocument] = []
    for index, feature in enumerate(payload.get("features", [])):
        if not isinstance(feature, dict):
            continue
        properties = feature.get("properties") or {}
        geometry = feature.get("geometry") or {}
        if not isinstance(properties, dict):
            properties = {}
        species = _resolve_species_from_attrs(properties, None)
        observed_at = _normalize_date_value(properties.get("observed_at") or properties.get("date") or properties.get("datum"))
        zone = _infer_zone(properties)
        geometry_type = str(geometry.get("type") or "unknown")
        title = species or path.stem
        text = _compose_feature_text(path, properties, species, observed_at, zone, geometry_type)
        docs.append(
            _build_document(
                project_id=project_id,
                project_slug=project_slug,
                source_root=source_root,
                path=path,
                category="geodata",
                title=title,
                text=text,
                species=species,
                observed_at=observed_at,
                zone=zone,
                geometry_type=geometry_type,
                metadata={
                    "row_index": index,
                    "feature_type": "geojson",
                    "attribute_keys": sorted(str(key) for key in properties.keys()),
                },
            )
        )

    return docs


def _scan_document_file(project_id: str, project_slug: str, source_root: Path, path: Path) -> list[RetrievalDocument]:
    text = _extract_text_from_file(path)
    if not text.strip():
        return [_build_file_document(project_id, project_slug, source_root, path, "document")]

    chunks = _chunk_text(text, max_chars=2000)
    documents = []
    for chunk_index, chunk in enumerate(chunks):
        documents.append(
            _build_document(
                project_id=project_id,
                project_slug=project_slug,
                source_root=source_root,
                path=path,
                category="document",
                title=path.stem,
                text=chunk,
                metadata={
                    "chunk_index": chunk_index,
                    "chunk_count": len(chunks),
                    "source_type": path.suffix.lower().lstrip("."),
                },
            )
        )
    return documents


def _build_file_document(project_id: str, project_slug: str, source_root: Path, path: Path, category: str) -> RetrievalDocument:
    title = path.stem
    text = f"{path.name} ({category})"
    return _build_document(
        project_id=project_id,
        project_slug=project_slug,
        source_root=source_root,
        path=path,
        category=category,
        title=title,
        text=text,
        metadata={
            "source_type": path.suffix.lower().lstrip("."),
            "file_only": True,
        },
    )


def _build_document(
    *,
    project_id: str,
    project_slug: str,
    source_root: Path,
    path: Path,
    category: str,
    title: str,
    text: str,
    species: str | None = None,
    observed_at: str | None = None,
    zone: str | None = None,
    geometry_type: str | None = None,
    metadata: dict[str, Any] | None = None,
) -> RetrievalDocument:
    relative_path = str(path.relative_to(source_root)).replace("\\", "/") if path.is_relative_to(source_root) else path.name
    return RetrievalDocument(
        id=_stable_document_id(project_id, path, title, category, species, observed_at, zone, geometry_type),
        project_id=project_id,
        project_slug=project_slug,
        source_path=str(path),
        relative_path=relative_path,
        file_name=path.name,
        extension=path.suffix.lower().lstrip("."),
        category=category,
        title=title,
        text=_clean_text(text),
        species=species,
        observed_at=observed_at,
        zone=zone,
        geometry_type=geometry_type,
        metadata=metadata or {},
        created_at=now_iso(),
    )


def _extract_text_from_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".csv", ".tsv", ".xml", ".html", ".htm", ".json", ".qgs", ".qml"}:
        try:
            return path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return ""
    if suffix == ".qgz":
        try:
            with zipfile.ZipFile(path) as archive:
                candidates = [name for name in archive.namelist() if name.lower().endswith(".qgs")]
                if candidates:
                    return archive.read(candidates[0]).decode("utf-8", errors="ignore")
        except Exception:
            return ""
        return ""
    if suffix == ".docx" and DocxDocument is not None:  # pragma: no cover - optional dependency
        try:
            document = DocxDocument(str(path))
            return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        except Exception:
            return ""
    if suffix == ".pdf" and PdfReader is not None:  # pragma: no cover - optional dependency
        try:
            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages[:20])
        except Exception:
            return ""
    return ""


def _compose_feature_text(
    path: Path,
    attrs: dict[str, Any],
    species: str | None,
    observed_at: str | None,
    zone: str | None,
    geometry_type: str | None,
) -> str:
    parts = [f"Datei: {path.name}"]
    if species:
        parts.append(f"Art: {species}")
    if observed_at:
        parts.append(f"Datum: {observed_at}")
    if zone:
        parts.append(f"Zone: {zone}")
    if geometry_type:
        parts.append(f"Geometrie: {geometry_type}")
    for key in sorted(attrs):
        value = attrs[key]
        if value in (None, ""):
            continue
        parts.append(f"{key}: {value}")
    return " | ".join(parts)


def _resolve_species_from_attrs(attrs: dict[str, Any], species_column: str | None) -> str | None:
    if species_column:
        value = attrs.get(species_column)
        if value not in (None, ""):
            resolved = infer_species_from_text(str(value))
            return resolved or str(value).strip()

    candidate_keys = [
        "species",
        "art",
        "artname",
        "taxon",
        "taxon_name",
        "wissenschaftlicher_name",
        "deutscher_name",
        "objektart",
        "bezeichnung",
        "name",
    ]
    for key in candidate_keys:
        value = attrs.get(key)
        if value in (None, ""):
            continue
        resolved = infer_species_from_text(str(value))
        if resolved:
            return resolved
    for key, value in attrs.items():
        if value in (None, ""):
            continue
        resolved = infer_species_from_text(str(value))
        if resolved:
            return resolved
    return None


def _infer_zone(attrs: dict[str, Any]) -> str | None:
    zone_keys = ("zone", "area", "gebiet", "district", "region", "ort", "place", "landscape", "flur")
    for key, value in attrs.items():
        key_normalized = normalize_species_name(str(key))
        if any(token in key_normalized for token in zone_keys) and value not in (None, ""):
            return str(value).strip()
    for key in zone_keys:
        if key in attrs and attrs[key] not in (None, ""):
            return str(attrs[key]).strip()
    return None


def _normalize_date_value(value: Any) -> str | None:
    if value in (None, ""):
        return None
    text = str(value).strip()
    if not text:
        return None
    try:
        parsed = datetime.fromisoformat(text.replace("Z", "+00:00"))
        return parsed.date().isoformat()
    except ValueError:
        pass
    try:
        parsed = datetime.strptime(text[:10], "%Y-%m-%d")
        return parsed.date().isoformat()
    except ValueError:
        return text


def _matches_filter(document: RetrievalDocument, filters: RetrievalFilter) -> bool:
    if filters.species:
        needle = normalize_species_name(_resolve_filter_species(filters.species) or filters.species)
        haystack = normalize_species_name(document.species or "")
        if needle not in haystack and haystack not in needle:
            return False
    if filters.file_type:
        if document.extension.lower() != filters.file_type.lower().lstrip("."):
            return False
    if filters.category:
        if document.category.casefold() != filters.category.casefold():
            return False
    if filters.zone:
        needle = normalize_species_name(filters.zone)
        haystack = normalize_species_name(document.zone or "")
        if needle not in haystack and haystack not in needle:
            return False
    if filters.date_from or filters.date_to:
        if document.observed_at is None:
            return False
        observed = _parse_date(document.observed_at)
        if observed is None:
            return False
        if filters.date_from:
            start = _parse_date(filters.date_from)
            if start is not None and observed < start:
                return False
        if filters.date_to:
            end = _parse_date(filters.date_to)
            if end is not None and observed > end:
                return False
    return True


def _parse_date(value: str | None) -> date | None:
    if not value:
        return None
    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00")).date()
    except ValueError:
        try:
            return datetime.strptime(value[:10], "%Y-%m-%d").date()
        except ValueError:
            return None


def _cosine_similarity(left: list[float], right: list[float]) -> float:
    if not left or not right:
        return 0.0
    if len(left) != len(right):
        right = right[: len(left)] + [0.0] * max(0, len(left) - len(right))
    dot = sum(a * b for a, b in zip(left, right))
    left_norm = math.sqrt(sum(value * value for value in left))
    right_norm = math.sqrt(sum(value * value for value in right))
    if left_norm == 0.0 or right_norm == 0.0:
        return 0.0
    return dot / (left_norm * right_norm)


def _lexical_boost(query: str, document: RetrievalDocument) -> float:
    if not query.strip():
        return 0.0
    query_tokens = _token_set(query)
    if not query_tokens:
        return 0.0
    doc_tokens = _token_set(" ".join([document.title, document.text, document.species or "", document.zone or ""]))
    overlap = len(query_tokens & doc_tokens)
    boost = overlap * 0.15
    if document.species and _token_set(document.species) & query_tokens:
        boost += 0.3
    if document.zone and _token_set(document.zone) & query_tokens:
        boost += 0.15
    return boost


def _make_snippet(text: str, query: str, max_length: int = 280) -> str:
    clean_text = _clean_text(text)
    if not clean_text:
        return ""
    if not query.strip():
        return clean_text[:max_length]
    query_tokens = [token for token in _token_set(query) if len(token) >= 3]
    if not query_tokens:
        return clean_text[:max_length]
    lowered = clean_text.casefold()
    best_index = None
    for token in query_tokens:
        index = lowered.find(token)
        if index >= 0 and (best_index is None or index < best_index):
            best_index = index
    if best_index is None:
        return clean_text[:max_length]
    start = max(0, best_index - 80)
    end = min(len(clean_text), best_index + max_length - 80)
    return clean_text[start:end].strip()


def _clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _tokenize(text: str) -> list[str]:
    cleaned = normalize_species_name(text)
    tokens = re.split(r"[^a-z0-9]+", cleaned)
    stopwords = {
        "und",
        "oder",
        "der",
        "die",
        "das",
        "the",
        "and",
        "fur",
        "für",
        "von",
        "mit",
        "auf",
        "in",
        "im",
        "of",
        "data",
        "file",
        "project",
        "projekt",
        "species",
        "art",
        "nachweis",
        "nachweise",
    }
    return [token for token in tokens if len(token) >= 2 and token not in stopwords]


def _token_set(text: str) -> set[str]:
    return set(_tokenize(text))


def _stable_document_id(
    project_id: str,
    path: Path,
    title: str,
    category: str,
    species: str | None,
    observed_at: str | None,
    zone: str | None,
    geometry_type: str | None,
) -> str:
    payload = "|".join(
        [
            project_id,
            str(path),
            title,
            category,
            species or "",
            observed_at or "",
            zone or "",
            geometry_type or "",
        ]
    )
    return hashlib.sha1(payload.encode("utf-8")).hexdigest()


def _find_column(columns: Iterable[Any], candidates: list[str]) -> str | None:
    lowered = {str(column).casefold(): str(column) for column in columns}
    for candidate in candidates:
        if candidate.casefold() in lowered:
            return lowered[candidate.casefold()]
    return None


def _qdrant_url() -> str | None:
    from os import getenv

    return getenv("HIPPO_AI_QDRANT_URL") or None


def _qdrant_path(index_root: Path) -> str | None:
    from os import getenv

    return getenv("HIPPO_AI_QDRANT_PATH") or str(index_root / "qdrant")


def _qdrant_filter(project_id: str, filters: RetrievalFilter):
    if not QDRANT_AVAILABLE:  # pragma: no cover - optional dependency
        return None

    conditions = [
        FieldCondition(key="project_id", match=MatchValue(value=project_id)),
    ]
    if filters.species:
        conditions.append(FieldCondition(key="species", match=MatchValue(value=_resolve_filter_species(filters.species) or filters.species)))
    if filters.file_type:
        conditions.append(FieldCondition(key="extension", match=MatchValue(value=filters.file_type.lstrip("."))))
    if filters.category:
        conditions.append(FieldCondition(key="category", match=MatchValue(value=filters.category)))
    if filters.zone:
        conditions.append(FieldCondition(key="zone", match=MatchValue(value=filters.zone)))
    if filters.date_from or filters.date_to:
        range_kwargs: dict[str, Any] = {}
        if filters.date_from:
            range_kwargs["gte"] = filters.date_from
        if filters.date_to:
            range_kwargs["lte"] = filters.date_to
        conditions.append(FieldCondition(key="observed_at", range=Range(**range_kwargs)))
    return Filter(must=conditions)


def _load_qdrant_backend(metadata: dict[str, Any], project_id: str) -> QdrantProjectIndexStore | None:
    if not QDRANT_AVAILABLE:
        return None
    backend = metadata.get("backend", "local")
    if backend != "qdrant":
        return None
    collection_name = metadata.get("qdrant_collection") or project_id
    try:
        return QdrantProjectIndexStore(
            project_id=project_id,
            collection_name=collection_name,
            url=metadata.get("qdrant_url") or _qdrant_url(),
            path=metadata.get("qdrant_path") or _qdrant_path(Path(metadata.get("index_path", "")) if metadata.get("index_path") else Path("workspace") / "state" / "retrieval"),
        )
    except Exception:
        return None


def _resolve_filter_species(value: str | None) -> str | None:
    if not value:
        return None
    return resolve_species_label(value) or value.strip()


def _document_from_payload(payload: dict[str, Any]) -> RetrievalDocument:
    data = dict(payload)
    data.pop("score", None)
    data.pop("id", None)
    return RetrievalDocument(**data)


def to_dict(summary: RetrievalIndexSummary | RetrievalSearchSummary) -> dict[str, Any]:
    payload = dataclasses.asdict(summary)
    if isinstance(summary, RetrievalSearchSummary):
        payload["hits"] = [dataclasses.asdict(hit) for hit in summary.hits]
    return payload
